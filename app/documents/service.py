import logging
from pathlib import Path

import aiofiles
from fastapi import UploadFile
from sqlalchemy.orm import Session

from app.config import settings
from app.documents.models import Document
from app.exceptions import bad_request, not_found
from app.investments.service import get_investment

logger = logging.getLogger(__name__)

CHUNK_SIZE = 64 * 1024  # 64 KB


def _validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise bad_request(
            f"File type '{ext}' not allowed. Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )
    return ext


def _investment_upload_dir(investment_name: str, security_id: int | None = None) -> Path:
    folder = settings.UPLOAD_ROOT / "investments" / investment_name
    if security_id is not None:
        folder = folder / f"security_{security_id}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _extract_text(file_path: Path) -> str:
    """Extract text content from PDF, Word, or Excel documents."""
    ext = file_path.suffix.lower()
    try:
        if ext == ".pdf":
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in (".docx", ".doc"):
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        elif ext == ".xlsx":
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" ".join(cells))
            wb.close()
            return "\n".join(parts)
        elif ext == ".xls":
            import xlrd

            wb = xlrd.open_workbook(file_path)
            parts = []
            for sheet in wb.sheets():
                for rx in range(sheet.nrows):
                    cells = [str(c.value) for c in sheet.row(rx) if c.value]
                    if cells:
                        parts.append(" ".join(cells))
            return "\n".join(parts)
    except Exception as e:
        logger.warning("Failed to extract text from %s: %s", file_path, e)
    return ""


async def _save_file(upload: UploadFile, dest: Path) -> int:
    total = 0
    async with aiofiles.open(dest, "wb") as f:
        while chunk := await upload.read(CHUNK_SIZE):
            total += len(chunk)
            if total > settings.MAX_FILE_SIZE:
                await f.close()
                dest.unlink(missing_ok=True)
                raise bad_request(
                    f"File exceeds maximum size of {settings.MAX_FILE_SIZE // (1024 * 1024)} MB"
                )
            await f.write(chunk)
    return total


def _validate_security(db: Session, investment_id: int, security_id: int | None):
    """Validate that the security belongs to the investment."""
    if security_id is None:
        return
    from app.securities.service import get_security
    get_security(db, investment_id, security_id)


async def upload_document(
    db: Session,
    investment_id: int,
    file: UploadFile,
    document_name: str,
    document_date: str | None = None,
    security_id: int | None = None,
) -> Document:
    investment = get_investment(db, investment_id)
    _validate_security(db, investment_id, security_id)
    original_filename = file.filename or "unknown"
    ext = _validate_extension(original_filename)

    folder = _investment_upload_dir(investment.investment_name, security_id)
    dest = folder / original_filename

    # Avoid overwriting — append counter if needed
    counter = 1
    while dest.exists():
        stem = Path(original_filename).stem
        dest = folder / f"{stem}_{counter}{ext}"
        counter += 1

    file_size = await _save_file(file, dest)

    # Extract text for search indexing
    text_content = _extract_text(dest)

    doc = Document(
        investment_id=investment.id,
        security_id=security_id,
        document_name=document_name,
        document_date=document_date,
        investment_series=investment.series,
        document_type=ext,
        file_path=str(dest),
        file_size=file_size,
        original_filename=original_filename,
        text_content=text_content,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


async def bulk_upload_documents(
    db: Session,
    investment_id: int,
    files: list[UploadFile],
    document_name: str,
    document_date: str | None = None,
    security_id: int | None = None,
) -> list[Document]:
    docs = []
    for file in files:
        doc = await upload_document(
            db, investment_id, file, document_name, document_date, security_id=security_id
        )
        docs.append(doc)
    return docs


def list_documents(
    db: Session, investment_id: int, security_id: int | None = None
) -> tuple[list[Document], int]:
    get_investment(db, investment_id)
    query = db.query(Document).filter(Document.investment_id == investment_id)
    if security_id is not None:
        query = query.filter(Document.security_id == security_id)
    docs = query.order_by(Document.created_at.desc()).all()
    return docs, len(docs)


def list_all_documents(db: Session) -> list[dict]:
    """List all documents across all investments with investment name, parse status, and workflow status."""
    from app.investments.models import Investment
    from app.financial_parsing.models import ParseJob
    from app.financial_parsing.workflow_service import compute_document_workflow

    docs = (
        db.query(Document, Investment.investment_name)
        .join(Investment, Document.investment_id == Investment.id)
        .order_by(Document.created_at.desc())
        .all()
    )

    # Get latest parse job status for each document
    parse_jobs = db.query(ParseJob).all()
    job_by_doc = {}
    for job in parse_jobs:
        existing = job_by_doc.get(job.document_id)
        if existing is None or job.created_at > existing.created_at:
            job_by_doc[job.document_id] = job

    result = []
    for doc, inv_name in docs:
        job = job_by_doc.get(doc.id)
        # Compute workflow status for PDF documents
        workflow = None
        if doc.document_type and doc.document_type.lower() == ".pdf":
            workflow = compute_document_workflow(db, doc.id)
        result.append({
            "id": doc.id,
            "investment_id": doc.investment_id,
            "investment_name": inv_name,
            "security_id": doc.security_id,
            "document_name": doc.document_name,
            "document_date": doc.document_date,
            "document_type": doc.document_type,
            "file_size": doc.file_size,
            "original_filename": doc.original_filename,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
            "parse_status": job.status if job else None,
            "workflow_status": workflow["workflow_status"] if workflow else None,
            "statement_count": workflow["statement_count"] if workflow else None,
            "mapped_count": workflow["mapped_count"] if workflow else None,
            "reviewed_count": workflow["reviewed_count"] if workflow else None,
            "approved_count": workflow["approved_count"] if workflow else None,
        })
    return result


def get_document(db: Session, investment_id: int, document_id: int) -> Document:
    doc = (
        db.query(Document)
        .filter(Document.id == document_id, Document.investment_id == investment_id)
        .first()
    )
    if not doc:
        raise not_found(f"Document {document_id} not found for investment {investment_id}")
    return doc


def delete_document(db: Session, investment_id: int, document_id: int) -> None:
    doc = get_document(db, investment_id, document_id)
    path = Path(doc.file_path)
    if path.exists():
        path.unlink()
    db.delete(doc)
    db.commit()
